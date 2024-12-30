import json
import os
import torch
import warnings
import re
from deepmultilingualpunctuation import PunctuationModel
from datetime import datetime
import tempfile
import spacy
import zipfile
from transformers import AutoTokenizer, AutoModelForTokenClassification
from docx import Document
from modules import moz_split as sp 
import csv
from modules import tab1_func as t1
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="transformers")
from multiprocessing import Process, Queue
import gradio as gr
import zipfile



# 警告を無視
warnings.filterwarnings("ignore", category=UserWarning, message="grouped_entities is deprecated and will be removed in version v5.0.0")
if torch.cuda.is_available():
    torch.cuda.empty_cache()
# デバイスの設定
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

def protect_inner_dots(text):
    # 置換ルールをリストに格納
    rules = [
        (r'(?<![\d\s])\.(?![\d\s]|$)', '[dot]'),  # ルール1: 数字以外の非空白文字 + "." + 数字以外の非空白文字 (文末は除外)
        (r'(?<![\d\s])\.(?=\d)', '[dot]'),        # ルール2: 数字以外の非空白文字 + "." + 数字
        (r'(?<=\d)\.(?![\d\s]|$)', '[dot]'),       # ルール3: 数字 + "." + 数字以外の非空白文字 (文末は除外)
        (r'(?<=\d)\.(?=\d)', '[dot]')               #　ルール4: 結局、小数点も含めてしまった。
    ]
    
    # ルールに従って置換
    for pattern, replacement in rules:
        text = re.sub(pattern, replacement, text)
      
    return text

def protect_special_cases_srt(text):
 
    with open("dot_manager.csv", newline='',encoding='utf-8') as dot_csvfile:
        reader = csv.reader(dot_csvfile)
        next(reader)
        dot_replacements = [(row[0],row[1]) for row in reader]
    
    
    for dot_original, dot_replacement in dot_replacements:
        r_dot_original=re.escape(dot_original)
        dot_new_original = rf"\b{r_dot_original}"
        text = re.sub(dot_new_original, dot_replacement, text)    
        
    text = protect_inner_dots(text)
    return text


def protect_special_cases_json(text):
   
    with open("dot_manager.csv", newline='',encoding='utf-8') as dot_csvfile:
        reader = csv.reader(dot_csvfile)
        next(reader)
        dot_replacements = [(row[0],row[1]) for row in reader]
    
    
    for dot_original, dot_replacement in dot_replacements:
        r_dot_original=re.escape(dot_original)
        dot_new_original = rf"\b{r_dot_original}"
        text = re.sub(dot_new_original, dot_replacement, text)

    text=protect_inner_dots(text)

    return text.lower()  # 小文字に変換


def restore_special_cases(text):
    text = text.replace('[dot]', '.')
    return text

def apply_custom_replacements(text):

    with open("replacements.csv", newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        next(reader)  # ヘッダーをスキップ
        replacements = [(row[0], row[1]) for row in reader]
    for original, replacement in replacements:
        original=re.escape(original)
        new_original = rf"\b{original}\b"
        text = re.sub(new_original, replacement, text)

    return text

'''def capitalize_sentences(text):
    sentences = text.split('.')
    capitalized_sentences = [sentence.strip().capitalize() + '.' for sentence in sentences if sentence]
    return ' '.join(capitalized_sentences)'''

def add_punctuation_and_transform(queue,texts):

    model_name = "oliverguhr/fullstop-punctuation-multilang-large"

    # PunctuationModelのインスタンスを作成
    punctuation_model = PunctuationModel(model=model_name)

    punctuated_texts = []
    total_steps = len(texts)

    #print(texts)
    for i, text in enumerate(texts):
        # PunctuationModelで句読点を復元
        if len(text)>0:

            punctuated_text = punctuation_model.restore_punctuation(text)
            
            
        else:
            punctuated_text = text
        punctuated_texts.append(punctuated_text)
        if i%10==0:
            queue.put(("progress",0.7*(i+1)/total_steps))


    #capitalized_texts = [capitalize_sentences(text) for text in punctuated_texts]

    # ":"と"- "を","に変換
    #transformed_texts = [text.replace(':', ',').replace('- ', ', ') for text in capitalized_texts]

    return punctuated_texts


def parse_srt(content):
    # 正規表現を使用して各セクションを抽出
    pattern = re.compile(r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n((?:.|\n)+?)(?=\n\d+\n|\Z)', re.MULTILINE)
    matches = pattern.finditer(content)
    subs = []

    for match in matches:
        index = int(match.group(1))
        start_time = match.group(2)
        end_time = match.group(3)
        text = match.group(4).strip()
        subs.append({'index': index, 'start': start_time, 'end': end_time, 'text': text})

    return subs

def process_srt_file(queue,input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()
    subs = parse_srt(content)
    
    texts = [sub['text'] for sub in subs]
    protected_texts = [protect_special_cases_srt(text) for text in texts]
    processed_texts = add_punctuation_and_transform(queue,protected_texts)
    
    for sub, processed_text in zip(subs, processed_texts):
        sub["text"] = processed_text
    
    with open(output_file, 'w', encoding='utf-8') as file:
        for sub in subs:
            file.write(f"{sub['index']}\n")
            file.write(f"{sub['start']} --> {sub['end']}\n")
            file.write(f"{sub['text']}\n\n")

def json_data_combine(data):
    combined_data = []
    i = 0

    while i < len(data):
        current_item = data[i]
        word = current_item['word']
        
        while (i + 1 < len(data)) and (not data[i + 1]['word'].startswith(" ")):
            next_item = data[i + 1]
            word += next_item['word']
            current_item['end'] = next_item['end']
            i += 1
        
        current_item['word'] = word
        combined_data.append(current_item)
        i += 1

    return combined_data

def clean_word(word):
    # [dot]全体を保護し、他の不要な文字を削除
    cleaned_word = re.sub(r"(?!\[dot\])[^\w\s'\[\]-]+", '', word).strip().lower()
    return cleaned_word.strip("-")

def load_json(json_file):
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            data = json_data_combine(data)
        for item in data:
            item['start'] = round(item['start'], 2)
            item['end'] = round(item['end'], 2)
            word = protect_special_cases_json(item['word'].strip())
            item['word'] = clean_word(word) # アポストロフィーを残す
            #print(f"JSON word: {item['word']}")  # デバッグ用出力
        return data
    except Exception as e:
        print(f"Error loading JSON file {json_file}: {e}")
        return []
        
def load_srt(srt_file):
    with open(srt_file, 'r', encoding='utf-8') as file:
        content = file.read()
    subs = parse_srt(content)
    return subs


def extract_key_words_with_context(sentence, previous_sentence=None):
    words = sentence.split()
    context_words = []

    # 前のセグメントから補完
    if previous_sentence:
        context_words = previous_sentence.split() + words
    else:
        context_words = words

    # 最後の5語を取得し、クリーン化
    combined_words = context_words[-5:]

    return tuple(clean_word(word) for word in combined_words)

def find_best_match(json_segment_data, key_word, segment_start_time):
    found_indices = []
    for i in range(len(json_segment_data) - len(key_word) + 1):
        json_phrase = tuple(json_segment_data[i + j]['word'] for j in range(len(key_word)))
        if key_word == json_phrase:
            found_indices.append(i)
    if found_indices:
        valid_matches = [(idx, json_segment_data[idx + len(key_word) - 1]['end']) for idx in found_indices if json_segment_data[idx + len(key_word) - 1]['end'] > segment_start_time]
        if valid_matches:
            best_match = min(valid_matches, key=lambda x: abs(x[1] - segment_start_time))
            return best_match[1]
    return None


def convert_to_seconds(subrip_time):
    if isinstance(subrip_time, str):
        h, m, s_ms = subrip_time.split(':')
        s, ms = s_ms.split(',')
        return int(h) * 3600 + int(m) * 60 + int(s) + int(ms) / 1000.0
    return round(subrip_time.hours * 3600 + subrip_time.minutes * 60 + subrip_time.seconds + subrip_time.milliseconds / 1000.0, 2)

def format_time(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = seconds % 60
    milliseconds = int((seconds - int(seconds)) * 1000)
    return f"{hours:02}:{minutes:02}:{int(seconds):02},{milliseconds:03}"


def split_srt_segment(segment, json_data,i):
    segment_start_time = convert_to_seconds(segment['start'])
    original_end_time = convert_to_seconds(segment['end'])
    
    #print(f"Segment start: {segment_start_time}, end: {original_end_time}, text: {segment['text']}")

    json_segment_data = [item for item in json_data if segment_start_time <= item['start'] <= original_end_time]
    #for item in json_segment_data:
       # print(f"JSON start: {item['start']}, end: {item['end']}, word: {item['word']}")

    protected_text = segment['text']
    words = protected_text.split()  # 単語単位で分割
    new_segments = []
    stock_sentences = []

    sentence = []  # 現在の文を保持するリスト

    for i, word in enumerate(words):
        sentence.append(word)
        
        # 単語の最後が ".", "!", "?" で終わるかを確認
        if word.endswith('.') or word.endswith('?') or word.endswith('!'):
            # sentenceの内容を連結して一つの文にする
            sentence_text = ' '.join(sentence)
            previous_sentence = sentence_text if stock_sentences else None  # 前の文は未使用に変更
            key_word = extract_key_words_with_context(sentence_text, previous_sentence)
            #print(f"Key words: {key_word}")  # デバッグ用出力

            if key_word:
                best_end_time = find_best_match(json_segment_data, key_word, segment_start_time)
                #print(f"Best end time: {best_end_time}")  # デバッグ用出力
            else:
                best_end_time = None

            if best_end_time is not None:
                if stock_sentences:
                    sentence_text = ' '.join(stock_sentences) + ' ' + sentence_text
                    stock_sentences = []
                segment_end_time = best_end_time
                new_segments.append({
                    'start': segment_start_time,
                    'end': segment_end_time,
                    'text': sentence_text
                })
                segment_start_time = segment_end_time
            elif i == len(words) - 1:
                if stock_sentences:
                    sentence_text = ' '.join(stock_sentences) + ' ' + sentence_text
                segment_end_time = original_end_time
                new_segments.append({
                    'start': segment_start_time,
                    'end': segment_end_time,
                    'text': sentence_text
                })
            else:
                stock_sentences.append(sentence_text)

            sentence = []  # 文のリセット

    if stock_sentences:
        segment_end_time = original_end_time
        new_segments.append({
            'start': segment_start_time,
            'end': segment_end_time,
            'text': ' '.join(stock_sentences)
        })

    #print(f"New segments for the current segment: {new_segments}")

    

    return new_segments


def process_segments(srt_subs, json_data):
    all_new_segments = []

    for i, segment in enumerate(srt_subs):
        #print(f"Processing segment: {segment}")
        new_segments = split_srt_segment(segment, json_data,i)
        all_new_segments.extend(new_segments)
    
    all_new_segments = fix_timestamp_inconsistencies(all_new_segments)
    return all_new_segments



def fix_timestamp_inconsistencies(segments):
    for i in range(1, len(segments)):
        if segments[i]['start'] < segments[i-1]['end']:
            segments[i]['start'] = segments[i-1]['end']
    return segments

def write_srt_file(segments, output_file):
    with open(output_file, 'w', encoding='utf-8') as f:
        for i, segment in enumerate(segments):
            start_time = format_time(segment['start'])
            end_time = format_time(segment['end'])
            text = restore_special_cases(segment['text'])
            text = apply_custom_replacements(text)
            f.write(f"{i+1}\n")
            f.write(f"{start_time} --> {end_time}\n")
            f.write(f"{text}\n\n")

def write_txt_file(segments, output_file):
    with open(output_file, 'w', encoding='utf-8') as f:
        full_text = " ".join(segment['text'] for segment in segments)
        full_text = restore_special_cases(full_text)
        full_text = apply_custom_replacements(full_text)
        f.write(full_text)


def tab9_main(queue,json_file, input_srt_file):
    # プログレスバーを表示
    try:
        queue.put(("print",f"{os.path.basename(json_file)}"))
        queue.put(("print",f"{os.path.basename(input_srt_file)}"))
        # 前処理
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
        os.makedirs(temp_dir, exist_ok=True)


        temp_srt_file = os.path.join(temp_dir,"temp_punctuated.srt")
        queue.put(("print","ピリオド付加を始めます。"))
        queue.put(("progress",0))
        process_srt_file(queue,input_srt_file, temp_srt_file)
        
        # JSONファイルと前処理したSRTファイルを読み込む
        json_data = load_json(json_file)
        srt_subs = load_srt(temp_srt_file)
        
        # 新しいセグメントを生成
        new_segments = process_segments(srt_subs, json_data)
        queue.put(("print","ピリオドを付け終わりました。"))
        queue.put(("progress",0.7))
        queue.put(("print","新たなピリオドに基づき字幕を分割しています。"))
        # 元のファイル名に_revを追加したファイル名を生成


        base_name = os.path.splitext(os.path.basename(input_srt_file))[0]
        temp_srt_output_file = os.path.join(temp_dir, f"deepmulti_{base_name}_rv.srt")
        temp_txt_output_file = os.path.join(temp_dir, f"deepmulti_{base_name}_rv_NR.txt")
        

        write_srt_file(new_segments, temp_srt_output_file)
        write_txt_file(new_segments, temp_txt_output_file)
        output_name1 =f'{base_name}_rv.srt'
        queue.put(("progress",0.8))
        queue.put(("print","固有名詞や文頭の大文字化を始めます。"))
        
        srt_output_file = sp.process_srt_file(queue,temp_srt_output_file,output_name1)
        output_name2=f'{base_name}_rv_NR.txt'
    
        txt_output_file,txtR_output_file = sp.process_text_file(temp_txt_output_file,output_name2)
        queue.put(("print","処理を終了しました。"))
        queue.put(("progress",1))
    

        return srt_output_file, txt_output_file,txtR_output_file
    except Exception as e:
        print(f"Error in tab9_main with files {json_file} and {input_srt_file}: {e}")
        return None, None, None






def repair(queue,json_files, srt_files):
    #print("JSON file path:", json_files)
    #print("SRT file path:", srt_files)
    valid_sets = []
    


    # JSONファイルとSRTファイルをファイル名で一致させてセットを作成
    for json_file in json_files:
        json_base_name = os.path.splitext(os.path.basename(json_file))[0]
        for srt_file in srt_files:
            srt_base_name = os.path.splitext(os.path.basename(srt_file))[0]
            if srt_base_name == json_base_name:
                valid_sets.append((json_file, srt_file))
                break
        
    if valid_sets==[]:

        return None,None

    srt_output_files = []
    txt_output_files = []
    txtR_output_files= []
    XLSX_output_files= []

    # プログレスバーの固定コンテナを作成
  
    total_sets = len(valid_sets)



    # 各セットを順次処理
    for i, (json_file, srt_file) in enumerate(valid_sets, start=1):
        
        try:

            srt_output_file, txt_output_file,txtR_output_file = tab9_main(
                queue,
                json_file, 
                srt_file,  
            )
            srt_output_files.append(srt_output_file)
            txt_output_files.append(txt_output_file)
            txtR_output_files.append(txtR_output_file)
            with open(srt_output_file,"r",encoding='utf-8') as f:
                srt_content = f.read()
            filename=os.path.splitext(os.path.basename(srt_output_file))[0]
            XLSX_file_path,_=t1.create_excel_from_srt_c(srt_content,filename)
            XLSX_output_files.append(XLSX_file_path)

        except Exception as e:
            message=f"Error processing file set {json_file} and {srt_file}: {e}"
            queue.put("error",message)
    
    if len(srt_output_files) > 1 :
        # ZIPファイルを生成
        timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
        os.makedirs(temp_dir, exist_ok=True)

        reversal_srt_zip = os.path.join(temp_dir, "reversal_srt_files.zip")
        reversal_txtNR_zip = os.path.join(temp_dir, "reversal_NR_files.zip")
        reversal_txtR_zip = os.path.join(temp_dir, "reversal_R_files.zip")
        reversal_XLSX_zip = os.path.join(temp_dir, "reversal_XLSX_files.zip")

        with zipfile.ZipFile(reversal_srt_zip, 'w') as srt_zip:
            for file in srt_output_files:
                srt_zip.write(file, os.path.basename(file))
        with zipfile.ZipFile(reversal_txtNR_zip, 'w') as txt_zip:
            for file in txt_output_files:
                txt_zip.write(file, os.path.basename(file))
        with zipfile.ZipFile(reversal_txtR_zip, 'w') as txt_zip:
            for file in txtR_output_files:
                txt_zip.write(file, os.path.basename(file))
        with zipfile.ZipFile(reversal_XLSX_zip, 'w') as xlsx_zip:
            for file in XLSX_output_files:
                xlsx_zip.write(file, os.path.basename(file))

        queue.put(("result", [reversal_srt_zip, reversal_txtNR_zip,reversal_txtR_zip,reversal_XLSX_zip]))
        queue.put(("done",None))
        return
    elif len(srt_output_files)==0:
 
        queue.put(("result",[None,None,None,None]))
        queue.put(("done",None))
        return
    # 単数ファイルの場合
    else:
        #print(srt_output_files[0])
        #print(txt_output_files[0])
        queue.put(("result",[srt_output_files[0], txt_output_files[0],txtR_output_files[0],XLSX_output_files[0]]))
        queue.put(("done",None))
        return

def run_spacy(json_files, srt_files,progress=gr.Progress()):
    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True)

    queue = Queue()
    process = Process(target=repair, args=(queue,json_files, srt_files))
    process.start()

    paths = None
    flag=False
 
    while True:
        message_type, data = queue.get()
        if message_type == "progress":
            progress(data)  # 進捗更新
        elif message_type == "result":
            paths = data  # パスのリストを受信
        elif message_type == "error":
            print(data)
            flag=True
            break 
        elif message_type == "print":
            print(data)
        elif message_type == "done":
            break
    process.join()
    if flag==True:
        return []
    
    return paths

def run_spacy2(main_files,progress=gr.Progress()):

    json_path=main_files[0]
    srt_path=main_files[1]
    

    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True)

    queue = Queue()
    process = Process(target=repair, args=(queue,[json_path],[srt_path]))
    process.start()

    paths = None
    flag=False
 
    while True:
        message_type, data = queue.get()
        if message_type == "progress":
            progress(data)  # 進捗更新
        elif message_type == "result":
            paths = data  # パスのリストを受信
        elif message_type == "error":
            print(data)
            flag=True
            break 
        elif message_type == "print":
            print(data)
        elif message_type == "done":
            break
    process.join()
    if flag==True:
        return []
    segment_info="ファイル修復が終わりました。"
    srt_output_path = paths[0]
    txt_nr_output_path = paths[1]
    txt_r_output_path = paths[2]
    XLSX_output_path = paths[3]
    with open(srt_output_path,"r",encoding="utf-8") as f:
        srt_content=f.read()
    with open(txt_nr_output_path,"r",encoding="utf-8") as f:
        txt_nr_content=f.read()
    with open(txt_r_output_path,"r",encoding="utf-8") as f:
        txt_r_content=f.read()
    _, df_display = t1.create_excel_from_srt_c(srt_content=srt_content, input_file_name="dammy")

    df_display=t1.dataframe_to_html_table(df_display)
    df_display=f"""
        <div class="my-table-container">
            {df_display}
        </div>
    """
    input_file_name=os.path.splitext(os.path.basename(srt_path))[0]
    html_srt = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{srt_content}</pre>"""
    html_nr_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_nr_content}</pre>"""
    html_r_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_r_content}</pre>"""
    filename_copy = os.path.splitext(os.path.basename(srt_output_path))[0]
    srt_dummy_output_path = srt_output_path


    # srtファイルからワードファイルへ変換




    ## txt(nr)をdoc変換
    txtdoc_nr = Document()
    txtdoc_nr_output_file_name = f"{input_file_name}_rv_txtnr.docx"
    txtdoc_nr_output_path = os.path.join(temp_dir, txtdoc_nr_output_file_name)

    with open(txt_nr_output_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    for line in lines:
        txtdoc_nr.add_paragraph(line)

    txtdoc_nr.save(txtdoc_nr_output_path)

    ## txt(r)をdoc変換
    txtdoc_r = Document()
    txtdoc_r_output_file_name = f"{input_file_name}_rv_txtr.docx"
    txtdoc_r_output_path = os.path.join(temp_dir, txtdoc_r_output_file_name)

    with open(txt_r_output_path, 'r', encoding='utf-8') as file:
        content = file.read()

    paragraph = txtdoc_r.add_paragraph()
    paragraph.add_run(content)  # 改行はそのまま出力
    txtdoc_r.save(txtdoc_r_output_path)

    # zipファイルにまとめる(srt,txtr,txtnr)。
    zip_core_file_name = f"{input_file_name}_rv_core.zip"
    zip_core_file_path = os.path.join(temp_dir, zip_core_file_name)

    with zipfile.ZipFile(zip_core_file_path, 'w') as zip_file:
        zip_file.write(json_path, os.path.basename(json_path))
        zip_file.write(srt_path,os.path.basename(srt_path))
        zip_file.write(srt_output_path, os.path.basename(srt_output_path))
        zip_file.write(txt_r_output_path, os.path.basename(txt_r_output_path))
        zip_file.write(txt_nr_output_path, os.path.basename(txt_nr_output_path))
        
    
    # zipファイルにまとめる(doc)。
    zip_doc_file_name = f"{input_file_name}_office_rv_en.zip"
    zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

    with zipfile.ZipFile(zip_doc_file_path, 'w') as zip_file:
        zip_file.write(XLSX_output_path, os.path.basename(XLSX_output_path))
        zip_file.write(txtdoc_nr_output_path, os.path.basename(txtdoc_nr_output_path))
        zip_file.write(txtdoc_r_output_path, os.path.basename(txtdoc_r_output_path))

    
    
    main_files = [
        srt_path,
        json_path,
        srt_output_path,
        txt_nr_output_path,
        txt_r_output_path,
        zip_core_file_path
    ]
    
    zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

    doc_files = [XLSX_output_path, txtdoc_nr_output_path, txtdoc_r_output_path, zip_doc_file_path]
    return segment_info,srt_content,txt_nr_content, txt_r_content, main_files, doc_files ,html_srt, html_nr_txt, html_r_txt, filename_copy, srt_dummy_output_path, df_display

